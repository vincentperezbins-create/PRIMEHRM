from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "manual_output" / "PRIMEHR_Complete_User_Manual_DepEd.docx"

BLUE = "1849A9"
DARK = "101828"
MUTED = "667085"
LIGHT_BLUE = "EFF4FF"
LIGHT_GRAY = "F9FAFB"
BORDER = "D0D5DD"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def normalize_table_geometry(table, widths_in: list[float]) -> None:
    widths_dxa = [int(round(width * 1440)) for width in widths_in]
    total_dxa = sum(widths_dxa)

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    old_grid = tbl.tblGrid
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(1, grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            cell.width = Inches(widths_in[idx])


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def add_header_footer(section) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.text = "PRIMEHR Complete User Manual"
    p.style = "Header"
    p.paragraph_format.space_after = Pt(4)
    p.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.name = "Arial"

    footer = section.footer
    p = footer.paragraphs[0]
    add_page_number(p)
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(MUTED)
        run.font.name = "Arial"


def add_paragraph(doc: Document, text: str, style: str = "Body Text"):
    p = doc.add_paragraph(text, style=style)
    return p


def add_note(doc: Document, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    normalize_table_geometry(table, [6.5])
    mark_header_row(table.rows[0])
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    set_cell_border(cell, "B2CCFF")
    set_cell_margins(cell, 120, 160, 120, 160)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)
    r.font.name = "Arial"
    r.font.size = Pt(10.5)
    p.add_run(" " + text)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    normalize_table_geometry(table, widths)
    mark_header_row(table.rows[0])
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        cell.width = Inches(widths[i])
        set_cell_width(cell, int(round(widths[i] * 1440)))
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_border(cell)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor.from_string(DARK)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            cells[i].width = Inches(widths[i])
            set_cell_width(cells[i], int(round(widths[i] * 1440)))
            set_cell_border(cells[i])
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9.2)
                    run.font.color.rgb = RGBColor.from_string(DARK)
    normalize_table_geometry(table, widths)
    doc.add_paragraph()


def numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        p.paragraph_format.space_after = Pt(4)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.space_after = Pt(4)


def configure_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    add_header_footer(section)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].font.color.rgb = RGBColor.from_string(DARK)
    styles["Body Text"].font.name = "Arial"
    styles["Body Text"].font.size = Pt(10.5)
    styles["Body Text"].paragraph_format.space_after = Pt(6)
    styles["Body Text"].paragraph_format.line_spacing = 1.08
    styles["Title"].font.name = "Arial"
    styles["Title"].font.size = Pt(24)
    styles["Title"].font.bold = True
    styles["Title"].font.color.rgb = RGBColor.from_string(BLUE)
    for name, size, color in (("Heading 1", 16, BLUE), ("Heading 2", 13, "344054"), ("Heading 3", 11.5, DARK)):
        style = styles[name]
        style.font.name = "Arial"
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)
    for list_name in ("List Number", "List Bullet"):
        style = styles[list_name]
        style.font.name = "Arial"
        style.font.size = Pt(10)
        style.paragraph_format.space_after = Pt(4)
    return doc


def cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph("Department of Education", style="Body Text")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(13)
    p.runs[0].font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph("Schools Division Office 1 Pangasinan", style="Body Text")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    doc.add_paragraph()
    p = doc.add_paragraph("PRIMEHR", style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("Complete User Manual", style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(20)
    p.runs[0].font.color.rgb = RGBColor.from_string(DARK)

    p = doc.add_paragraph("Human Resource Management System", style="Body Text")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(13)
    p.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    for _ in range(5):
        doc.add_paragraph()
    add_table(
        doc,
        ["Document", "Prepared For", "Coverage", "Date"],
        [[
            "User Manual",
            "DepEd office personnel, school personnel, validators, and system administrators",
            "Dashboard, Employees, Leave Management, Attendance, Reports, and User Roles",
            date.today().strftime("%B %d, %Y"),
        ]],
        [1.15, 2.15, 2.15, 1.05],
    )
    doc.add_page_break()


def front_matter(doc: Document) -> None:
    doc.add_heading("Document Control", level=1)
    add_table(
        doc,
        ["Item", "Details"],
        [
            ["System Name", "PRIMEHR Human Resource Management System"],
            ["Office Use", "For DepEd internal office, school, and division HR workflows."],
            ["Primary Users", "Admin, Employee, School Head, District Supervisor, Chief, Unit Head, and authorized validators."],
            ["Confidentiality", "Contains operational guidance only. Do not include account passwords, private employee documents, or personally sensitive reports in distributed copies."],
        ],
        [1.5, 5.0],
    )

    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. System Overview",
        "2. Login Instructions",
        "3. Navigation Guide",
        "4. Dashboard",
        "5. Employees",
        "6. Leave Management",
        "7. Attendance",
        "8. Reports",
        "9. User Roles and Permissions",
        "10. Admin Functions",
        "11. Troubleshooting",
        "12. Frequently Asked Questions",
        "13. Glossary",
    ]
    bullets(doc, toc_items)
    doc.add_page_break()


def system_overview(doc: Document) -> None:
    doc.add_heading("1. System Overview", level=1)
    add_paragraph(doc, "PRIMEHR is an internal human resource management system used to organize personnel profiles, 201 file submissions, leave records, performance form submissions, reports, user roles, and related HR administrative workflows.")
    add_table(
        doc,
        ["Module", "Purpose", "Common Users"],
        [
            ["Dashboard", "Shows summaries, counts, status cards, and quick links for each user type.", "All users"],
            ["Employees", "Maintains employee accounts, assigned school or office, role, and 201 file records.", "Admin, School Head, validators"],
            ["Leave Management", "Processes leave applications, balances, ledgers, transactions, and approved leave reporting.", "Employees, Admin, School Head, authorized validators"],
            ["Attendance", "Supports attendance-related monitoring in L&D participant workflows and leave-related workforce tracking.", "Admin, L&D managers, School Head, validators"],
            ["Reports", "Exports filtered lists and approved leave data using Excel, PDF, CSV, and print tools where available.", "Admin and authorized reporting users"],
            ["User Roles", "Controls access to menus, validations, school-level records, and administrative functions.", "Admin"],
        ],
        [1.35, 3.35, 1.8],
    )
    add_note(doc, "Office reminder:", "Menu visibility depends on the account role and validator assignments configured by the administrator.")


def login(doc: Document) -> None:
    doc.add_heading("2. Login Instructions", level=1)
    doc.add_heading("2.1 Before Logging In", level=2)
    bullets(doc, [
        "Use a supported web browser such as Google Chrome, Microsoft Edge, or Mozilla Firefox.",
        "Confirm that the computer is connected to the local or hosted PRIMEHR address supplied by the ICT or system administrator.",
        "Prepare the username or email address and password issued by the office.",
    ])
    doc.add_heading("2.2 Sign In", level=2)
    numbered(doc, [
        "Open the PRIMEHR login page.",
        "Enter the assigned username or email address in the username field.",
        "Enter the password exactly as provided. Passwords are case-sensitive.",
        "Click Login.",
        "After successful login, review the landing page and sidebar menu available to your role.",
    ])
    doc.add_heading("2.3 Sign Out", level=2)
    numbered(doc, [
        "Click the user profile area or logout option in the top navigation.",
        "Select Logout.",
        "Close the browser tab if the computer is shared or located in a public office area.",
    ])
    add_note(doc, "Account security:", "Never share passwords, leave accounts open on shared workstations, or allow another employee to process records using your account.")


def navigation(doc: Document) -> None:
    doc.add_heading("3. Navigation Guide", level=1)
    add_paragraph(doc, "PRIMEHR uses a left sidebar for module navigation, a top navigation bar for user and notification controls, and data tables for searching, filtering, viewing, and exporting records.")
    add_table(
        doc,
        ["Area", "What It Does", "User Action"],
        [
            ["Left Sidebar", "Groups menus such as Home, Self Service, Administration, Records, HR Workflows, Performance, Growth, and Profile.", "Click a menu title to expand it, then choose a subpage."],
            ["Top Bar", "Shows user controls, notifications, and page-level actions where available.", "Use it for profile access, alerts, and logout."],
            ["Data Tables", "Display lists such as users, leave applications, schools, districts, document types, ledgers, and reports.", "Use search, filters, pagination, export, print, view, update, approve, or return actions."],
            ["Forms", "Capture employee details, leave requests, performance submissions, certificates, and setup records.", "Complete required fields marked by the page, review entries, then submit."],
        ],
        [1.25, 3.25, 2.0],
    )
    doc.add_heading("3.1 General Navigation Procedure", level=2)
    numbered(doc, [
        "Log in to PRIMEHR.",
        "Open the required sidebar menu.",
        "Use filters or search boxes to narrow records.",
        "Open View, Add, Update, Approve, Return, Export, or Print depending on the task.",
        "Read confirmation messages before leaving the page.",
    ])


def dashboard(doc: Document) -> None:
    doc.add_heading("4. Dashboard", level=1)
    add_paragraph(doc, "Dashboards provide a quick operational picture. The exact cards and links vary by role.")
    add_table(
        doc,
        ["Dashboard", "Typical Content", "Recommended Use"],
        [
            ["Admin Dashboard", "System-wide user, school, 201 file, leave, and activity summaries.", "Use at the start of the day to identify pending validations and data maintenance needs."],
            ["Leave Credits Dashboard", "Leave credit summaries and balance-related indicators.", "Use before processing leave balances or reviewing ledger issues."],
            ["School Dashboard", "School-level employee records and leave monitoring.", "Use to monitor personnel assigned to the school."],
            ["My Dashboard", "Personal 201 file status, leave balance, and recent user activity.", "Use by employees to check personal compliance and pending submissions."],
            ["Rewards and L&D Dashboards", "Recognition, training, participant, attendance, evaluation, and certificate indicators.", "Use to monitor employee development activities."],
        ],
        [1.55, 3.0, 1.95],
    )
    doc.add_heading("4.1 Review Dashboard Items", level=2)
    numbered(doc, [
        "Open Home from the sidebar.",
        "Choose the dashboard available to your account.",
        "Review summary cards, counts, and recent activity areas.",
        "Click the relevant linked module when a count or status requires action.",
        "Refresh the page after completing related records to confirm the updated count.",
    ])


def employees(doc: Document) -> None:
    doc.add_heading("5. Employees", level=1)
    add_paragraph(doc, "The Employees area includes account management, employee profile data, role assignments, school or office assignment, and 201 file document tracking.")
    doc.add_heading("5.1 Add a User Account", level=2)
    numbered(doc, [
        "Log in as an administrator.",
        "Open Settings, then User List.",
        "Click Add or Add User.",
        "Enter the employee name, email or username, password, role, personnel type, school or office assignment, and other required profile details.",
        "Assign validator permissions only when the user is officially authorized to validate 201 files, OPCRF, IPCRF, or leave.",
        "Save the record and confirm that the user appears in the list.",
    ])
    doc.add_heading("5.2 Update an Employee Account", level=2)
    numbered(doc, [
        "Open Settings, then User List.",
        "Search for the employee by name, employee number, school, office, or role.",
        "Click Update or the available edit action.",
        "Correct the profile, school, division unit, office unit, role, office role, or validator assignment.",
        "Save changes and ask the employee to log out and log in again if menu access changed.",
    ])
    doc.add_heading("5.3 Upload Employee 201 Files", level=2)
    numbered(doc, [
        "Log in using the employee account.",
        "Open My Submission, then My 201 Files.",
        "Select the required document type.",
        "Attach a clear scanned copy or accepted file format.",
        "Submit the file and monitor the status for approval or returned remarks.",
    ])
    doc.add_heading("5.4 Validate 201 Files", level=2)
    numbered(doc, [
        "Log in as Admin or an authorized 201 validator.",
        "Open 201 Files, then Validate 201 Files.",
        "Filter or search for the employee record.",
        "Open the uploaded document and check completeness, readability, and correct document type.",
        "Approve if acceptable, or return with specific remarks if correction is needed.",
    ])


def leave_management(doc: Document) -> None:
    doc.add_heading("6. Leave Management", level=1)
    add_paragraph(doc, "Leave Management handles leave applications, approved leave reports, balances, ledger entries, transactions, and leave type setup.")
    doc.add_heading("6.1 Apply for Leave", level=2)
    numbered(doc, [
        "Open My Leave, then Apply Leave.",
        "Select the leave type.",
        "Enter the inclusive dates and number of days or schedule details requested by the form.",
        "Provide purpose, remarks, or supporting information.",
        "Review the balance warning or eligibility message when the leave type is credit-based.",
        "Submit the application and monitor Leave History for status updates.",
    ])
    doc.add_heading("6.2 Review Leave Applications", level=2)
    numbered(doc, [
        "Open Leave Management, then Leave Applications. School Heads open School Leave, then Leave Applications.",
        "Use filters for employee, school, office, leave type, month, year, or status.",
        "Open the application details.",
        "Check the applicant, dates, leave type, available credits, remarks, and attachments when required.",
        "Approve, reject, return, or update status according to office policy and your assigned permission.",
    ])
    doc.add_heading("6.3 Manage Leave Balances", level=2)
    numbered(doc, [
        "Open Leave Management, then Leave Balances.",
        "Search for the employee.",
        "Review vacation leave, sick leave, CTO, or other balances shown by the system.",
        "Use the adjust balance function only for authorized corrections.",
        "Record clear remarks for every manual adjustment.",
    ])
    doc.add_heading("6.4 Review Leave Ledger", level=2)
    numbered(doc, [
        "Open Leave Management, then Leave Ledger.",
        "Search for the employee and review chronological balance movements.",
        "Confirm that approved applications and transactions appear in the ledger.",
        "Investigate missing, duplicate, or incorrect entries before issuing official reports.",
    ])
    doc.add_heading("6.5 Maintain Leave Types and Transactions", level=2)
    numbered(doc, [
        "Open Leave Management, then Leave Types to maintain the list of leave categories.",
        "Open Transactions to review or encode balance-related movements.",
        "Check whether the leave type affects credits and pay status before saving.",
        "Coordinate with HR records staff before changing leave type configuration used by active employees.",
    ])


def attendance(doc: Document) -> None:
    doc.add_heading("7. Attendance", level=1)
    add_paragraph(doc, "In the current PRIMEHR build, attendance is primarily handled through Learning and Development participant monitoring, attendance responses, certificate eligibility, and leave-related availability records. It is not shown as a separate daily time record sidebar module.")
    doc.add_heading("7.1 Monitor Training Attendance", level=2)
    numbered(doc, [
        "Open Learning and Development, then All Trainings or Participants.",
        "Select the training record.",
        "Review nominated participants, attendance responses, evaluation status, and certificate readiness.",
        "Generate certificates only after verifying attendance and completion requirements.",
    ])
    doc.add_heading("7.2 Validate Attendance-Related Records", level=2)
    numbered(doc, [
        "Confirm that the participant is assigned to the correct training or activity.",
        "Check attendance response completeness and related evaluation entries.",
        "Resolve duplicate or mismatched participant details before issuing certificates.",
        "Use reports or exports for official monitoring files when needed.",
    ])
    doc.add_heading("7.3 Leave and Attendance Coordination", level=2)
    bullets(doc, [
        "Approved leave records should be checked when preparing workforce availability reports.",
        "School and office heads should review pending leave before confirming staffing schedules.",
        "Any separate biometric or DTR system should remain the official source unless PRIMEHR is formally configured for daily attendance capture.",
    ])


def reports(doc: Document) -> None:
    doc.add_heading("8. Reports", level=1)
    add_paragraph(doc, "Reports are available through filtered tables and export buttons on selected pages. Common export formats include Excel, PDF, CSV, and Print depending on the module.")
    add_table(
        doc,
        ["Report Area", "Where to Open", "Typical Filters"],
        [
            ["User List", "Settings > User List", "Role, personnel type, school, division unit, office unit, keyword search"],
            ["Approved Leave Applications", "Leave Management > Approved Leave Report", "Employee ID, name, school, office, year, month"],
            ["Leave Balances", "Leave Management > Leave Balances", "Employee, school, office, balance category"],
            ["Leave Ledger", "Leave Management > Leave Ledger", "Employee, year, leave type, transaction period"],
            ["201 File Status", "201 Files > Validate 201 Files or School Employee 201 Files", "Employee, document type, status, school"],
            ["L&D Participants", "Learning and Development > Participants", "Training, school, participant status, attendance, certificate status"],
        ],
        [1.8, 2.45, 2.25],
    )
    doc.add_heading("8.1 Generate a Report", level=2)
    numbered(doc, [
        "Open the page containing the report data.",
        "Apply the required filters before exporting.",
        "Review the table results on screen.",
        "Click Export Excel, Export PDF, CSV, or Print if available.",
        "Check the downloaded or printed file before submitting it as an official attachment.",
    ])
    add_note(doc, "Reporting control:", "Reports reflect the records currently saved in PRIMEHR. Validate source records before using reports for official transmittals.")


def roles(doc: Document) -> None:
    doc.add_heading("9. User Roles and Permissions", level=1)
    add_table(
        doc,
        ["Role", "Role ID", "Typical Access"],
        [
            ["Admin", "1", "Full administration, user list, schools, districts, document types, notifications, audit logs, manuals, 201 validation, leave administration, performance forms, reports, L&D, and rewards."],
            ["Office Head", "2", "Office-related submissions, validation or management functions depending on assignment, rewards, L&D, and personal self-service."],
            ["School Head", "3", "School dashboard, school employee 201 files, school leave applications, school leave balances, school leave ledger, staff nomination, and personal self-service."],
            ["Employee", "4", "Personal dashboard, 201 file upload, leave application, leave history, leave balance, OPCRF/IPCRF submissions, recognitions, L&D, and profile."],
            ["District Supervisor", "5", "District-level user access plus authorized management, L&D, rewards, and validation functions as assigned."],
            ["Chief", "6", "Division chief user access plus authorized management, L&D, rewards, and validation functions as assigned."],
            ["Unit Head", "7", "Office or unit head access, OPCRF responsibilities, authorized management workflows, and personal self-service."],
        ],
        [1.35, 0.65, 4.5],
    )
    doc.add_heading("9.1 Assign Validator Permissions", level=2)
    numbered(doc, [
        "Open Settings, then User List.",
        "Add or update the target user account.",
        "Enable only the validation permissions required by the employee's official designation: 201, OPCRF, IPCRF, or Leave.",
        "Save the account.",
        "Ask the user to log out and log back in to refresh the sidebar menu.",
    ])


def admin_functions(doc: Document) -> None:
    doc.add_heading("10. Admin Functions", level=1)
    add_table(
        doc,
        ["Function", "Menu Path", "Admin Procedure"],
        [
            ["Manage Users", "Settings > User List", "Add, update, filter, export, print, and assign roles or validator permissions."],
            ["Manage Schools", "Settings > School List", "Add or update school records used in employee assignment and reports."],
            ["Manage Districts", "Settings > District List", "Maintain district reference data used by schools and users."],
            ["Manage Document Types", "Settings > Document Types / 201 Files", "Create and update accepted 201 file categories."],
            ["Send Notifications", "Settings > Notifications", "Create notifications for all users, a role, or a selected user."],
            ["Review Audit Logs", "Settings > Audit Logs", "Check user actions, exports, print activities, and administrative activity history."],
            ["Generate User Manuals", "Settings > User Manual Generator", "Download role-based manuals with screenshots and procedural notes."],
            ["Leave Setup", "Leave Management > Leave Types / Transactions / Adjust Balance", "Maintain leave categories, transactions, and authorized balance corrections."],
            ["Performance Setup", "Performance Forms > Offices / Units", "Manage OPCRF offices or units and review OPCRF/IPCRF submissions."],
        ],
        [1.55, 2.1, 2.85],
    )
    doc.add_heading("10.1 Admin Daily Checklist", level=2)
    bullets(doc, [
        "Review dashboard counts for pending 201 files, leave applications, and new records.",
        "Check notifications and audit logs for unusual or failed activities.",
        "Validate urgent employee submissions before office cut-off times.",
        "Export required reports only after filters and source records are confirmed.",
        "Keep role and validator assignments aligned with office orders or official designations.",
    ])


def troubleshooting(doc: Document) -> None:
    doc.add_heading("11. Troubleshooting", level=1)
    add_table(
        doc,
        ["Issue", "Likely Cause", "Resolution"],
        [
            ["Cannot log in", "Incorrect credentials, inactive account, or wrong URL.", "Re-enter credentials, verify Caps Lock, confirm PRIMEHR URL, then request password reset or account check from Admin."],
            ["Access denied", "User role does not allow the page.", "Ask Admin to confirm role ID and validator permissions."],
            ["Menu is missing", "Role or validator assignment does not include the module.", "Log out and log back in after Admin updates the account."],
            ["Upload fails", "File is too large, unsupported, corrupted, or network is unstable.", "Use a clear accepted file, reduce file size if needed, and try again on a stable connection."],
            ["Leave balance looks incorrect", "Missing ledger transaction, pending application, or manual adjustment issue.", "Check Leave Ledger and Transactions; escalate to HR/Admin before making corrections."],
            ["Report export is blank", "No results after filtering or export scripts did not load.", "Clear filters, refresh the page, and try the export again."],
            ["Print window blocked", "Browser pop-up blocker is active.", "Allow pop-ups for the PRIMEHR site and repeat Print."],
            ["Certificate not generated", "Attendance or evaluation completion is missing.", "Review participant attendance and evaluation records before generating certificates."],
        ],
        [1.65, 2.2, 2.65],
    )


def faq(doc: Document) -> None:
    doc.add_heading("12. Frequently Asked Questions", level=1)
    questions = [
        ("Who can create user accounts?", "Only Admin accounts can create and update user accounts."),
        ("Why do different users see different menus?", "PRIMEHR shows menus based on role ID and validator permissions."),
        ("Can an employee submit leave without available credits?", "Credit-based leave types should be checked against the employee balance. Follow HR policy for leave types that do not require credits."),
        ("How will I know if my 201 file was returned?", "Open My 201 Files and check the status or remarks. Returned files should be corrected and resubmitted."),
        ("Can School Heads view all division personnel?", "No. School Head access is normally limited to the assigned school and school-level workflows."),
        ("What should I do before exporting a report?", "Apply filters, review the on-screen data, and confirm that source records are correct."),
        ("Who can validate leave, OPCRF, IPCRF, or 201 files?", "Admin can validate by default. Other users need the corresponding validator permission assigned by Admin."),
        ("What if my profile information is wrong?", "Report the correction to Admin or authorized HR staff so the user record can be updated."),
    ]
    for q, a in questions:
        doc.add_heading(q, level=2)
        add_paragraph(doc, a)


def glossary(doc: Document) -> None:
    doc.add_heading("13. Glossary", level=1)
    add_table(
        doc,
        ["Term", "Meaning"],
        [
            ["201 File", "Official employee personnel file containing required employment documents."],
            ["Admin", "System user with full configuration and management access."],
            ["Audit Logs", "System records showing selected user actions and administrative activities."],
            ["CTO", "Compensatory Time Off, when enabled and tracked in leave credits."],
            ["Dashboard", "Summary page showing counts, status cards, and quick links."],
            ["Employee", "Standard user role for personal submissions, leave, performance forms, and profile access."],
            ["IPCRF", "Individual Performance Commitment and Review Form."],
            ["Leave Ledger", "Chronological record of leave credit additions, deductions, and adjustments."],
            ["Leave Type", "Configured category of leave used by the application form and ledger."],
            ["OPCRF", "Office Performance Commitment and Review Form."],
            ["Role ID", "Numeric role assignment used by PRIMEHR to control access."],
            ["Validator", "Authorized user who reviews and approves or returns submissions in assigned areas."],
        ],
        [1.5, 5.0],
    )


def build() -> None:
    doc = configure_document()
    cover(doc)
    front_matter(doc)
    system_overview(doc)
    login(doc)
    navigation(doc)
    dashboard(doc)
    employees(doc)
    leave_management(doc)
    attendance(doc)
    reports(doc)
    roles(doc)
    admin_functions(doc)
    troubleshooting(doc)
    faq(doc)
    glossary(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
