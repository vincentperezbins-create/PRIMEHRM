from __future__ import annotations

import json
import re
import zipfile
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "manual_output" / "PRIMEHR_System_Documentation_Package"
BLUE = "1849A9"
DARK = "101828"
MUTED = "667085"
LIGHT_BLUE = "EFF4FF"
BORDER = "D0D5DD"


MODULES = [
    ("Dashboard", "Dashboard cards and summaries for administrators, school heads, employees, L&D, rewards, and leave credits."),
    ("Profile Management", "Employee profile viewing and updates, including personal, contact, work assignment, and account information."),
    ("User Management", "Administrative account creation, role assignment, validator permissions, account status, and password maintenance."),
    ("Employee / 201 Files", "Employee document upload, document type management, validation, returned remarks, and school-level viewing."),
    ("Leave Management", "Leave application filing, review, approval, balance maintenance, ledger entries, transactions, and Form 6 generation."),
    ("Attendance", "Training attendance and participant monitoring through L&D workflows; leave records support staff availability review."),
    ("Reports", "Filtered exports, print actions, audit CSV, approved leave reports, user lists, school lists, and ledger tables."),
    ("Settings", "School, district, document type, notification, audit log, office/unit, role, and manual generation administration."),
    ("Performance Forms", "Office/Unit OPCRF, Employee IPCRF, indicators, MOVs, ratings, review statuses, and validation logs."),
    ("Rewards and Recognition", "Award programs, nominations, supporting documents, recognition certificates, and nominee monitoring."),
    ("Learning and Development", "Programs, trainings, participants, attendance/evaluation tracking, certificate submissions, and generated certificates."),
]

ROLE_ROWS = [
    ["Admin", "1", "Full system administration, users, settings, validation, reports, leave, OPCRF/IPCRF, rewards, L&D, and audit logs."],
    ["Office Head", "2", "Office-level user; may manage or validate assigned workflows depending on validator permissions."],
    ["School Head", "3", "School dashboard, school employees, school leave records, school 201 files, nominations, L&D participation."],
    ["Employee", "4", "Personal dashboard, profile, 201 file uploads, leave applications, IPCRF/OPCRF submissions, certificates."],
    ["District Supervisor", "5", "District-level or program-owner workflows plus assigned validation or L&D/reward functions."],
    ["Chief", "6", "Division chief-level workflows plus assigned management, validation, L&D, or rewards functions."],
    ["Unit Head", "7", "Unit-level OPCRF or office workflows plus assigned validation and self-service access."],
]

TABLE_PURPOSE = {
    "sdopang1_user": "Stores user accounts, employee profile data, roles, school/office assignment, and validator permissions.",
    "sdopang1_roles": "Reference table for application roles and access group names.",
    "sdopang1schoollist": "Stores school master list data used for employee assignment and school-level reports.",
    "sdopang1_documents": "Stores uploaded 201 file documents and validation status.",
    "sdopang1_document_types": "Reference list of accepted 201 file document categories.",
    "leave_applications": "Stores employee leave application requests and review details.",
    "leave_balances": "Stores leave credit balances per user and leave type.",
    "leave_transactions": "Stores leave credit additions, deductions, and adjustments.",
    "leave_types": "Reference table for leave type configuration.",
    "leave_year_start": "Stores yearly leave start or baseline data.",
    "system_runs": "Tracks scheduled leave accrual or entitlement processing runs.",
    "sdopang1_ipcrf": "Stores employee IPCRF submissions and validation status.",
    "sdopang1_opcrf": "Stores OPCRF records for offices or units.",
    "sdopang1_opcrf_indicators": "Stores OPCRF indicators, success indicators, and ratings.",
    "sdopang1_opcrf_movs": "Stores OPCRF means of verification files.",
    "sdopang1_opcrf_logs": "Stores OPCRF review or action history.",
    "sdopang1_offices": "Stores office/unit hierarchy and office head assignments.",
    "audit_logs": "Stores system action logs for monitoring, exports, logins, and administrative activity.",
    "sdopang1_notifications": "Stores notifications targeted to users, roles, or all users.",
    "ld_categories": "Stores L&D category reference values.",
    "ld_programs": "Stores L&D master programs.",
    "ld_trainings": "Stores L&D training records.",
    "ld_participants": "Stores L&D participant status, attendance, evaluation, and certificate information.",
    "ld_certificate_submissions": "Stores user-submitted external certificates for review.",
    "ld_generated_certificates": "Stores certificates generated from legacy training attendance records.",
    "reward_categories": "Stores reward and recognition category values.",
    "reward_programs": "Stores award programs and nomination windows.",
    "reward_nominations": "Stores submitted award nominations.",
    "reward_documents": "Stores files attached to reward nominations.",
    "reward_recognitions": "Stores final recognition awards and certificate paths.",
}


def split_sql_columns(body: str) -> list[str]:
    parts, current, depth, quote = [], [], 0, None
    for ch in body:
        if quote:
            current.append(ch)
            if ch == quote:
                quote = None
            continue
        if ch in "'\"`":
            quote = ch
            current.append(ch)
        elif ch == "(":
            depth += 1
            current.append(ch)
        elif ch == ")":
            depth -= 1
            current.append(ch)
        elif ch == "," and depth == 0:
            parts.append("".join(current).strip())
            current = []
        else:
            current.append(ch)
    if current:
        parts.append("".join(current).strip())
    return parts


def extract_schema() -> dict:
    sources = list((ROOT / "database").glob("*.sql")) + [
        ROOT / "accounts" / "core" / "audit.php",
        ROOT / "accounts" / "core" / "rewards_helpers.php",
        ROOT / "accounts" / "core" / "ld_helpers.php",
    ]
    text = "\n".join(p.read_text(errors="ignore") for p in sources if p.exists())
    create_re = re.compile(r"CREATE TABLE(?: IF NOT EXISTS)?\s+`?([A-Za-z0-9_]+)`?\s*\((.*?)\)\s*(?:ENGINE|;)", re.S | re.I)
    tables = {}
    for name, body in create_re.findall(text):
        fields = []
        primary, foreign = [], []
        indexes = []
        for item in split_sql_columns(body):
            line = re.sub(r"\s+", " ", item.strip()).rstrip(",")
            if not line:
                continue
            up = line.upper()
            if up.startswith("PRIMARY KEY"):
                primary += re.findall(r"`([^`]+)`", line)
                continue
            if "FOREIGN KEY" in up:
                m = re.search(r"FOREIGN KEY\s*\(`?([A-Za-z0-9_]+)`?\)\s*REFERENCES\s*`?([A-Za-z0-9_]+)`?\s*\(`?([A-Za-z0-9_]+)`?\)", line, re.I)
                if m:
                    foreign.append({"field": m.group(1), "references": f"{m.group(2)}.{m.group(3)}"})
                continue
            if up.startswith(("KEY ", "INDEX ", "UNIQUE KEY", "CONSTRAINT ")):
                indexes.append(line)
                continue
            m = re.match(r"`?([A-Za-z0-9_]+)`?\s+(.+)", line)
            if m:
                fields.append({"name": m.group(1), "type": m.group(2)})
                if "PRIMARY KEY" in up:
                    primary.append(m.group(1))
        tables[name] = {"fields": fields, "primary": sorted(set(primary)), "foreign": foreign, "indexes": indexes}

    for alter in re.finditer(r"ALTER TABLE\s+`?([A-Za-z0-9_]+)`?\s+(.*?);", text, re.S | re.I):
        table, body = alter.group(1), alter.group(2)
        if table not in tables:
            tables[table] = {"fields": [], "primary": [], "foreign": [], "indexes": []}
        if "ADD PRIMARY KEY" in body.upper():
            tables[table]["primary"] += [x for x in re.findall(r"`([^`]+)`", body) if x not in tables[table]["primary"]]
        for m in re.finditer(r"FOREIGN KEY\s*\(`?([A-Za-z0-9_]+)`?\)\s*REFERENCES\s*`?([A-Za-z0-9_]+)`?\s*\(`?([A-Za-z0-9_]+)`?\)", body, re.I):
            rel = {"field": m.group(1), "references": f"{m.group(2)}.{m.group(3)}"}
            if rel not in tables[table]["foreign"]:
                tables[table]["foreign"].append(rel)
    return dict(sorted(tables.items()))


def table_widths(widths):
    return [int(round(w * 1440)) for w in widths]


def set_cell(cell, text="", fill=None, bold=False):
    cell.text = text
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        node = borders.find(qn(f"w:{edge}")) or OxmlElement(f"w:{edge}")
        if node.getparent() is None:
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:color"), BORDER)
    if fill:
        shd = tc_pr.find(qn("w:shd")) or OxmlElement("w:shd")
        if shd.getparent() is None:
            tc_pr.append(shd)
        shd.set(qn("w:fill"), fill)
    mar = tc_pr.first_child_found_in("w:tcMar") or OxmlElement("w:tcMar")
    if mar.getparent() is None:
        tc_pr.append(mar)
    for m, v in (("top", 80), ("start", 120), ("bottom", 80), ("end", 120)):
        node = mar.find(qn(f"w:{m}")) or OxmlElement(f"w:{m}")
        if node.getparent() is None:
            mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs:
            r.font.name = "Arial"
            r.font.size = Pt(9)
            r.font.bold = bold
            r.font.color.rgb = RGBColor.from_string(DARK)


def normalize_table(table, widths):
    wd = table_widths(widths)
    total = sum(wd)
    table.autofit = False
    tbl = table._tbl
    pr = tbl.tblPr
    for child in list(pr):
        if child.tag in (qn("w:tblW"), qn("w:tblLayout")):
            pr.remove(child)
    tbl_w = OxmlElement("w:tblW")
    pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    layout = OxmlElement("w:tblLayout")
    pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    old = tbl.tblGrid
    if old is not None:
        tbl.remove(old)
    grid = OxmlElement("w:tblGrid")
    for w in wd:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    tbl.insert(1, grid)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths[i])
            cell_pr = cell._tc.get_or_add_tcPr()
            for child in list(cell_pr):
                if child.tag == qn("w:tcW"):
                    cell_pr.remove(child)
            tcw = OxmlElement("w:tcW")
            cell_pr.append(tcw)
            tcw.set(qn("w:w"), str(wd[i]))
            tcw.set(qn("w:type"), "dxa")


def mark_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = tr_pr.find(qn("w:tblHeader")) or OxmlElement("w:tblHeader")
    if node.getparent() is None:
        tr_pr.append(node)
    node.set(qn("w:val"), "true")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    normalize_table(table, widths)
    mark_header(table.rows[0])
    for i, h in enumerate(headers):
        set_cell(table.cell(0, i), h, LIGHT_BLUE, True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell(cells[i], str(val))
    normalize_table(table, widths)
    doc.add_paragraph()
    return table


def add_note(doc, title, body):
    add_table(doc, [title], [[body]], [6.5])


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.space_after = Pt(3)


def numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        p.paragraph_format.space_after = Pt(3)


def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    for typ, txt in (("begin", None), ("instrText", "PAGE"), ("separate", None), ("end", None)):
        if typ == "instrText":
            node = OxmlElement("w:instrText")
            node.set(qn("xml:space"), "preserve")
            node.text = txt
        else:
            node = OxmlElement("w:fldChar")
            node.set(qn("w:fldCharType"), typ)
        run._r.append(node)


def new_doc(title, subtitle):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(0.8)
    section.left_margin = section.right_margin = Inches(0.85)
    header = section.header.paragraphs[0]
    header.text = title
    header.runs[0].font.name = "Arial"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    page_number(section.footer.paragraphs[0])
    for st in doc.styles:
        if st.type == 1 and st.name in ("Normal", "Body Text", "List Bullet", "List Number"):
            st.font.name = "Arial"
            st.font.size = Pt(10.5)
    doc.styles["Title"].font.name = "Arial"
    doc.styles["Title"].font.size = Pt(22)
    doc.styles["Title"].font.bold = True
    doc.styles["Title"].font.color.rgb = RGBColor.from_string(BLUE)
    for name, size, color in (("Heading 1", 16, BLUE), ("Heading 2", 13, "344054"), ("Heading 3", 11, DARK)):
        s = doc.styles[name]
        s.font.name = "Arial"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(9)
        s.paragraph_format.space_after = Pt(5)
    doc.add_paragraph("Department of Education", style="Body Text").alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("Schools Division Office 1 Pangasinan", style="Body Text")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    p = doc.add_paragraph(title, style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(subtitle, style="Body Text")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    doc.add_paragraph()
    add_table(doc, ["Document", "System", "Date"], [[title, "PRIMEHR", date.today().strftime("%B %d, %Y")]], [2.6, 2.0, 1.9])
    doc.add_page_break()
    doc.add_heading("Table of Contents", 1)
    add_note(doc, "Note", "Use Microsoft Word References > Update Table after opening this file if automatic page-numbered TOC is required.")
    return doc


def save(doc, filename):
    path = OUT_DIR / filename
    doc.save(path)
    return path


def add_toc_list(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_page_break()


def system_overview(schema):
    doc = new_doc("System Overview", "Purpose, scope, users, features, and technology stack")
    add_toc_list(doc, ["System Title", "Purpose", "Target Users", "Main Features", "Scope and Objectives", "Technology Stack"])
    doc.add_heading("1. System Title", 1)
    doc.add_paragraph("PRIMEHR - Human Resource Management System for DepEd Schools Division Office 1 Pangasinan.", style="Body Text")
    doc.add_heading("2. Purpose of the System", 1)
    doc.add_paragraph("PRIMEHR centralizes personnel records, 201 file submissions, leave credits, leave applications, performance forms, learning and development, rewards, notifications, reports, and administrative monitoring.", style="Body Text")
    doc.add_heading("3. Target Users", 1)
    add_table(doc, ["User Group", "Purpose"], [[r[0], r[2]] for r in ROLE_ROWS], [1.7, 4.8])
    doc.add_heading("4. Main Features", 1)
    add_table(doc, ["Feature", "Description"], MODULES, [2.1, 4.4])
    doc.add_heading("5. Scope and Objectives", 1)
    bullets(doc, [
        "Provide authenticated access for HR, school, office, and employee users.",
        "Maintain employee master records and role-based access.",
        "Accept and validate required 201 file submissions.",
        "Process leave applications and maintain leave balances and ledgers.",
        "Support OPCRF/IPCRF performance documentation workflows.",
        "Generate reports and maintain audit logs for accountability.",
    ])
    doc.add_heading("6. Technology Stack Used", 1)
    add_table(doc, ["Layer", "Technology Detected", "Notes"], [
        ["Server-side", "PHP", "Page-based PHP application using includes, helpers, and PDO."],
        ["Database", "MySQL / MariaDB", "Database name in config is monitoring."],
        ["Database Access", "PDO", "PDO error mode exceptions, default associative fetch, emulated prepares disabled."],
        ["Frontend", "Bootstrap 4 / DeskApp", "Admin dashboard template with responsive sidebar and cards."],
        ["JavaScript", "jQuery, DataTables, SweetAlert, plugins", "Used for tables, filtering, exports, alerts, and UI controls."],
        ["Document Generation", "PHPWord", "Used for Microsoft Word manual generation."],
        ["QR / Forms / PDF", "phpqrcode, PDF/XLSX templates", "Used for QR support, leave Form 6, and L&D certificates."],
    ], [1.5, 2.0, 3.0])
    return save(doc, "01_SYSTEM_OVERVIEW.docx")


def user_manual(schema):
    doc = new_doc("User Manual", "Step-by-step guide for users and office personnel")
    add_toc_list(doc, ["Login", "Dashboard Navigation", "Profile Management", "User Management", "Employee Management", "Attendance", "Leave Management", "Reports", "Settings", "Logout"])
    sections = [
        ("Login", "Authenticate using registered DepEd email and password.", ["Email address", "Password"], ["Sign In button submits credentials.", "Invalid login alert displays when credentials are wrong."]),
        ("Dashboard Navigation", "Review summaries, quick links, counts, and recent activities.", ["Dashboard cards", "Sidebar menu", "Top bar notifications"], ["Open Home or Dashboard from the sidebar.", "Click linked cards to go to related records."]),
        ("Profile Management", "View or update personal and employment profile data.", ["Name", "Contact information", "School/office", "Role", "Position"], ["Save updates after checking required fields.", "Contact Admin for locked role or assignment fields."]),
        ("User Management", "Admin-only account management.", ["Employee details", "Email", "Password", "Role", "Validator permissions"], ["Add creates a new account.", "Update changes account details.", "Delete removes a user when permitted."]),
        ("Employee Management", "Manage employee records and uploaded 201 file documents.", ["Document type", "Year", "File attachment", "Status", "Remarks"], ["Upload submits the document.", "Approve accepts a valid document.", "Return sends it back with remarks."]),
        ("Attendance", "Monitor L&D attendance, participant completion, and certificate eligibility.", ["Training", "Participant", "Attendance status", "Evaluation status"], ["Generate certificates after attendance/evaluation checks.", "Use participant lists for monitoring reports."]),
        ("Leave Management", "Submit, approve, and monitor leave requests and credits.", ["Leave type", "Date range", "Purpose", "Balance", "Status"], ["Apply Leave files a request.", "Approve/Reject/Return handles review.", "Ledger shows credit movements."]),
        ("Reports Generation", "Export filtered lists and monitoring data.", ["Filters", "Search", "Export Excel/PDF/CSV", "Print"], ["Filter before export.", "Review on-screen results before printing or attaching."]),
        ("Settings", "Admin setup for schools, districts, document types, notifications, audit logs, offices, and manuals.", ["Master data", "Status", "Descriptions", "Assignments"], ["Add or update reference records carefully.", "Avoid deleting records used by active transactions."]),
        ("Logout", "End the authenticated session safely.", ["Logout link"], ["Click Logout.", "Close browser on shared computers."]),
    ]
    for idx, (title, purpose, fields, functions) in enumerate(sections, 1):
        doc.add_heading(f"{idx}. {title}", 1)
        doc.add_paragraph(purpose, style="Body Text")
        doc.add_paragraph("[Insert Screenshot Here]", style="Body Text")
        doc.add_heading("Buttons and Functions", 2)
        bullets(doc, functions)
        doc.add_heading("Forms and Fields", 2)
        bullets(doc, fields)
        doc.add_heading("Workflow", 2)
        numbers(doc, ["Open the module from the sidebar.", "Review the page title and available controls.", "Complete or filter the form fields.", "Click the appropriate action button.", "Read the success or error message before leaving the page."])
    return save(doc, "02_USER_MANUAL.docx")


def admin_guide(schema):
    doc = new_doc("Administrator Guide", "Responsibilities, permissions, maintenance, backup, and troubleshooting")
    add_toc_list(doc, ["Admin Responsibilities", "Users and Roles", "Permissions", "Database Backup", "Maintenance", "Security Reminders", "Common Issues"])
    doc.add_heading("1. Admin Responsibilities", 1)
    bullets(doc, ["Maintain accurate user accounts and roles.", "Monitor pending 201 files, leave requests, OPCRF/IPCRF submissions, notifications, and audit logs.", "Protect employee data and restrict access to authorized personnel only.", "Coordinate database backup and system maintenance schedules."])
    doc.add_heading("2. Managing Users and Roles", 1)
    add_table(doc, ["Role", "ID", "Access Summary"], ROLE_ROWS, [1.3, .6, 4.6])
    numbers(doc, ["Open Settings > User List.", "Add or update the employee account.", "Assign the correct role, school, division unit, office unit, and office role.", "Enable validator permissions only when officially authorized.", "Ask the user to log out and log back in after permission changes."])
    doc.add_heading("3. Managing Permissions", 1)
    bullets(doc, ["Admin role has broad access.", "Non-admin validators use can_validate_201, can_validate_opcrf, can_validate_ipcrf, and can_validate_leave fields.", "School Head role is scoped to school-level records where implemented.", "L&D and Rewards use role scope helpers to determine admin, school head, employee, or program-owner behavior."])
    doc.add_heading("4. Database Backup Procedures", 1)
    numbers(doc, ["Open phpMyAdmin from XAMPP.", "Select the monitoring database.", "Click Export.", "Choose Quick for routine backup or Custom for selected tables.", "Select SQL format.", "Save the exported file with date and time in the backup folder.", "Test restore in a separate local database before relying on a backup."])
    doc.add_heading("5. System Maintenance", 1)
    bullets(doc, ["Back up database and uploaded files before code updates.", "Review audit logs for unusual activities.", "Check upload folders for storage usage.", "Run leave accrual scripts only under approved HR schedule.", "Keep PHP, Apache, MySQL, and browser versions updated in the local environment."])
    doc.add_heading("6. Security Reminders", 1)
    bullets(doc, ["Use strong administrator passwords.", "Never share admin credentials.", "Avoid using root MySQL with blank password in production.", "Restrict upload file types and folder execution permissions.", "Review user roles regularly."])
    doc.add_heading("7. Troubleshooting Common Issues", 1)
    add_table(doc, ["Issue", "Resolution"], [["Access denied", "Confirm role and validator permissions."], ["Cannot export PDF", "Refresh page and check DataTables/pdfmake assets."], ["Upload fails", "Check file size, extension, and upload folder permissions."], ["Leave balance mismatch", "Review ledger and transactions before manual adjustment."]], [2.0, 4.5])
    return save(doc, "03_ADMINISTRATOR_GUIDE.docx")


def technical_doc(schema):
    doc = new_doc("Technical Documentation", "Architecture, files, configuration, sessions, CRUD, and libraries")
    add_toc_list(doc, ["Folder Structure", "File Structure", "PHP Architecture", "Database Flow", "Authentication", "CRUD", "Libraries", "Configuration"])
    doc.add_heading("1. Folder Structure Explanation", 1)
    rows = [
        ["accounts/", "Main authenticated PHP pages, modules, helpers, classes, partials, cron scripts, and uploads."],
        ["accounts/core/", "Authentication, CSRF, audit, leave, OPCRF, notifications, L&D, rewards, and form generation helper functions."],
        ["accounts/partials/", "Reusable page fragments: head, navbar, sidebar, footer, session, preloader."],
        ["accounts/classes/", "PHP class models such as User."],
        ["assets/ and assets_pang1/", "Custom UI, images, logos, and local visual assets."],
        ["database/", "SQL dumps, migrations, templates, and form files."],
        ["vendor/", "Composer packages, including PHPWord."],
        ["phpqrcode/", "QR code generation library."],
        ["manual_output/", "Generated manuals and documentation package outputs."],
    ]
    add_table(doc, ["Folder", "Purpose"], rows, [2.0, 4.5])
    doc.add_heading("2. File Structure Explanation", 1)
    add_table(doc, ["File Pattern", "Purpose"], [
        ["admin_*.php", "Admin pages for setup, validation, reports, and management."],
        ["user_*.php", "Employee self-service pages."],
        ["school_*.php", "School Head scoped pages."],
        ["admin_ajax_*.php", "DataTables or AJAX JSON endpoints."],
        ["admin_query_*.php", "Action handlers for create, update, delete, and validation operations."],
        ["*_helpers.php", "Domain helper functions and reusable SQL logic."],
    ], [2.0, 4.5])
    doc.add_heading("3. PHP Architecture Overview", 1)
    bullets(doc, ["Page-controller style PHP application.", "Each page includes core/db.php for database access and authentication helpers as needed.", "Shared visual layout comes from partial files.", "Business logic is split across helper files and query/action handlers.", "AJAX endpoints return JSON to DataTables and dynamic UI screens."])
    doc.add_heading("4. Database Connection Flow", 1)
    numbers(doc, ["config/config.php returns host, database, user, and password.", "core/db.php loads configuration.", "PDO DSN uses mysql:host plus database name and utf8mb4 charset.", "PDO options enable exceptions, associative fetches, and native prepared statements.", "A Database helper and User model are initialized."])
    doc.add_heading("5. Authentication and Session Handling", 1)
    bullets(doc, ["accounts/login.php starts the session and verifies email/password against sdopang1_user.", "password_verify is used for stored password hashes.", "session_regenerate_id(true) is called after login to reduce session fixation risk.", "require_login redirects unauthenticated users.", "require_role restricts pages to allowed role IDs.", "require_validator checks specific validator columns for non-admin users."])
    doc.add_heading("6. CRUD Operation Explanation", 1)
    add_table(doc, ["Operation", "Typical Pattern"], [["Create", "Prepared INSERT into module table, then redirect or JSON success response."], ["Read", "SELECT queries with joins and DataTables filters."], ["Update", "Prepared UPDATE through admin_query or update pages."], ["Delete", "Prepared DELETE for admin-only master data removal."], ["Upload", "Validate file input, move_uploaded_file, save file path and metadata."]], [1.3, 5.2])
    doc.add_heading("7. Libraries and Plugins Used", 1)
    bullets(doc, ["PHPWord for DOCX generation.", "phpqrcode for QR generation.", "Bootstrap 4 / DeskApp for layout.", "jQuery and DataTables for searchable tables.", "DataTables Buttons, pdfmake, JSZip, and print plugins for exports.", "SweetAlert and assorted UI plugins from the DeskApp template."])
    doc.add_heading("8. Important Configuration Files", 1)
    add_table(doc, ["File", "Purpose"], [["config/config.php", "Database credentials and database name."], ["core/db.php", "PDO connection and helper initialization."], ["accounts/core/auth.php", "Login, role, and validator access checks."], ["accounts/core/csrf.php", "CSRF token generation and verification."], ["accounts/partials/leftsidebar.php", "Role-based navigation menu."], ["composer.json", "PHP dependency declaration for PHPWord."], ["package.json", "Frontend build/template dependency metadata."]], [2.4, 4.1])
    return save(doc, "04_TECHNICAL_DOCUMENTATION.docx")


def database_doc(schema):
    doc = new_doc("Database Documentation", "Database dictionary, keys, and relationship explanations")
    add_toc_list(doc, ["Database Dictionary", "Detected Tables", "Relationships", "Relationship Explanations"])
    doc.add_heading("1. Database Dictionary Summary", 1)
    add_table(doc, ["Metric", "Value"], [["Detected tables", str(len(schema))], ["Primary database", "monitoring"], ["Schema sources", "database/*.sql, audit.php, rewards_helpers.php, ld_helpers.php"]], [2.0, 4.5])
    doc.add_heading("2. Detected Tables", 1)
    add_table(doc, ["Table Name", "Purpose", "Fields", "Primary Key"], [[name, TABLE_PURPOSE.get(name, "Application data table detected from SQL or runtime schema."), str(len(meta["fields"])), ", ".join(meta["primary"]) or "Not declared in source"] for name, meta in schema.items()], [1.9, 3.1, .7, .8])
    doc.add_heading("3. Database Dictionary", 1)
    for name, meta in schema.items():
        doc.add_heading(name, 2)
        doc.add_paragraph(TABLE_PURPOSE.get(name, "Application table detected from project source."), style="Body Text")
        rows = []
        fks = {x["field"]: x["references"] for x in meta["foreign"]}
        for f in meta["fields"]:
            rows.append([f["name"], f["type"], "Yes" if f["name"] in meta["primary"] else "", fks.get(f["name"], "")])
        add_table(doc, ["Field Name", "Data Type / Definition", "PK", "FK / Relationship"], rows or [["No fields parsed", "", "", ""]], [1.6, 3.1, .5, 1.3])
    doc.add_heading("4. Relationship Explanations", 1)
    rel_rows = []
    for name, meta in schema.items():
        for fk in meta["foreign"]:
            rel_rows.append([name, fk["field"], fk["references"], "Child record references parent record."])
    if not rel_rows:
        rel_rows = [["No explicit FK parsed", "", "", "Several relationships are implemented through joins even when SQL constraints are not declared."]]
    add_table(doc, ["Source Table", "Field", "References", "Explanation"], rel_rows, [1.6, 1.4, 1.7, 1.8])
    return save(doc, "05_DATABASE_DOCUMENTATION.docx")


def installation_doc(schema):
    doc = new_doc("Installation Guide", "XAMPP setup, database import, configuration, and local run instructions")
    add_toc_list(doc, ["Prerequisites", "XAMPP Setup", "Database Import", "Configuration", "PHP Extensions", "Permissions", "Run Locally"])
    doc.add_heading("1. Prerequisites", 1)
    bullets(doc, ["Windows workstation or server.", "XAMPP with Apache, MySQL/MariaDB, and PHP.", "Web browser.", "Project folder copied to htdocs.", "Database SQL dump or migration files."])
    doc.add_heading("2. XAMPP Setup", 1)
    numbers(doc, ["Install XAMPP.", "Start Apache and MySQL from the XAMPP Control Panel.", "Open http://localhost/phpmyadmin.", "Confirm PHP can access mysqli/PDO MySQL extensions."])
    doc.add_heading("3. Database Import Steps", 1)
    numbers(doc, ["Create a database named monitoring.", "Open phpMyAdmin > monitoring > Import.", "Import database/monitoring (1).sql if available because it contains the latest leave and OPCRF structures.", "Run migration SQL files if the dump is older.", "Confirm that the detected PRIMEHR tables exist."])
    doc.add_heading("4. Configuration Instructions", 1)
    add_table(doc, ["File", "Value"], [["config/config.php", "host=localhost, db=monitoring, user=root, pass='' for default local XAMPP."], ["core/db.php", "Uses config values to create PDO connection."], ["Browser URL", "http://localhost/PRIMEHR/accounts/login.php"]], [2.3, 4.2])
    doc.add_heading("5. Required PHP Extensions", 1)
    bullets(doc, ["pdo_mysql", "mbstring", "zip for DOCX/PHPWord operations", "fileinfo for upload validation when implemented", "gd if image processing or QR image output is required"])
    doc.add_heading("6. Folder Permissions", 1)
    bullets(doc, ["Allow write access to upload folders under accounts/uploads if used.", "Allow write access to manual_output for generated manuals.", "Protect SQL dumps and backup folders from public web access in production."])
    doc.add_heading("7. Running the Project Locally", 1)
    numbers(doc, ["Copy PRIMEHR into XAMPP htdocs.", "Start Apache and MySQL.", "Import database.", "Verify config/config.php.", "Open the login URL.", "Sign in using a valid account from sdopang1_user."])
    return save(doc, "06_INSTALLATION_GUIDE.docx")


def troubleshooting_doc(schema):
    doc = new_doc("Troubleshooting Guide", "Common system problems and solutions")
    add_toc_list(doc, ["Login Issues", "Database Issues", "Missing Files", "Permission Issues", "PHP Extension Issues", "Export/Print Issues"])
    issues = [
        ["Invalid email or password", "Wrong credentials or no account record.", "Check email spelling, reset password, verify account in sdopang1_user."],
        ["Inactive account", "User status is inactive.", "Admin must reactivate account if authorized."],
        ["Database connection failed", "Incorrect config or MySQL service stopped.", "Start MySQL and verify config/config.php."],
        ["Access denied", "Role or validator permission mismatch.", "Admin should verify role_id and can_validate columns."],
        ["Missing image or file", "File moved, bad path, or upload directory missing.", "Check file_path value and actual upload folder."],
        ["Upload fails", "Extension, size, or permission problem.", "Check folder permissions and accepted file type logic."],
        ["Export PDF does not work", "DataTables/pdfmake assets failed to load.", "Refresh page, check browser console, verify vendor scripts."],
        ["Blank report", "Filters exclude all data.", "Clear filters and test with a known record."],
        ["PHP extension error", "Required extension disabled.", "Enable pdo_mysql, zip, mbstring, gd, or fileinfo in php.ini and restart Apache."],
    ]
    add_table(doc, ["Issue", "Possible Cause", "Solution"], issues, [1.7, 2.2, 2.6])
    return save(doc, "07_TROUBLESHOOTING_GUIDE.docx")


def security_doc(schema):
    doc = new_doc("Security Recommendations", "Secure coding, database, session, upload, and role-based controls")
    add_toc_list(doc, ["SQL Injection", "Password Hashing", "Session Security", "Uploads", "RBAC", "Secure Coding"])
    doc.add_heading("1. SQL Injection Prevention", 1)
    bullets(doc, ["Continue using PDO prepared statements for all user-supplied values.", "Avoid interpolating table names, column names, or WHERE clauses from request data.", "Validate DataTables ordering and filter parameters against allowlists."])
    doc.add_heading("2. Password Hashing", 1)
    bullets(doc, ["Use password_hash with PASSWORD_DEFAULT for new or changed passwords.", "Use password_verify during login.", "Never store or display plain text passwords.", "Force password change after temporary password reset."])
    doc.add_heading("3. Session Security", 1)
    bullets(doc, ["Keep session_regenerate_id(true) after login.", "Set secure and httponly cookie flags in production HTTPS deployments.", "Expire sessions after inactivity.", "Destroy session on logout."])
    doc.add_heading("4. File Upload Validation", 1)
    bullets(doc, ["Validate file extension and MIME type.", "Limit file size.", "Store uploads outside executable folders when possible.", "Rename uploaded files safely.", "Scan or review sensitive employee documents before approval."])
    doc.add_heading("5. Role-Based Access Control", 1)
    bullets(doc, ["Use require_role for page-level access.", "Use require_validator for specific validation workflows.", "Review role assignments regularly.", "Record sensitive actions in audit_logs."])
    doc.add_heading("6. Secure Coding Practices", 1)
    bullets(doc, ["Escape output with htmlspecialchars.", "Use CSRF tokens on state-changing forms.", "Avoid exposing PHP errors to users in production.", "Remove or protect SQL dumps and cheat/test files before deployment.", "Use least-privilege MySQL credentials instead of root with blank password in production."])
    return save(doc, "08_SECURITY_RECOMMENDATIONS.docx")


def glossary_doc(schema):
    doc = new_doc("Glossary", "Technical and system terms used in PRIMEHR")
    add_toc_list(doc, ["Glossary Terms"])
    terms = [
        ["201 File", "Official personnel file containing required employee documents."],
        ["Admin", "System user with broad administrative access."],
        ["AJAX", "Browser request that updates page data without reloading the full page."],
        ["Audit Log", "Record of important user or system actions."],
        ["CRUD", "Create, Read, Update, Delete operations."],
        ["DataTables", "JavaScript table plugin used for searching, filtering, pagination, and export."],
        ["DepEd", "Department of Education."],
        ["IPCRF", "Individual Performance Commitment and Review Form."],
        ["L&D", "Learning and Development."],
        ["MySQL", "Relational database used by PRIMEHR."],
        ["OPCRF", "Office Performance Commitment and Review Form."],
        ["PDO", "PHP Data Objects database abstraction used for MySQL queries."],
        ["Role ID", "Numeric user access group stored in sdopang1_user.role_id."],
        ["Session", "Server-side login state for the current browser user."],
        ["Validator", "Authorized user who can approve or return submissions."],
    ]
    add_table(doc, ["Term", "Definition"], terms, [1.8, 4.7])
    return save(doc, "09_GLOSSARY.docx")


def appendices_doc(schema):
    doc = new_doc("Appendices", "Sample workflows, screenshot placeholders, and process notes")
    add_toc_list(doc, ["Sample Workflows", "Screenshot Placeholders", "System Process Notes"])
    doc.add_heading("1. Sample Workflows", 1)
    workflows = [
        ("Employee Leave Filing", ["Employee logs in.", "Employee opens My Leave > Apply Leave.", "Employee completes leave type, dates, and purpose.", "System records pending leave.", "Authorized reviewer approves, rejects, or returns application.", "Ledger and reports reflect approved records."]),
        ("201 File Validation", ["Employee uploads document.", "Validator opens Validate 201 Files.", "Validator checks file and document type.", "Validator approves or returns with remarks.", "Employee monitors status and resubmits if needed."]),
        ("Admin User Creation", ["Admin opens User List.", "Admin clicks Add User.", "Admin enters profile and role.", "Admin assigns school/office and validator permissions.", "User signs in and verifies sidebar access."]),
        ("L&D Certificate Generation", ["Admin or program owner opens L&D training.", "Participant attendance/evaluation records are reviewed.", "System generates certificates for eligible participants.", "Certificates are printed, downloaded, or reviewed."]),
    ]
    for title, steps in workflows:
        doc.add_heading(title, 2)
        numbers(doc, steps)
    doc.add_heading("2. Sample Screenshot Placeholders", 1)
    for m, _ in MODULES:
        doc.add_paragraph(f"{m}: [Insert Screenshot Here]", style="Body Text")
    doc.add_heading("3. System Process Notes", 1)
    bullets(doc, ["Menu access depends on role and validator permissions.", "Reports should be generated only after reviewing filters.", "Database backups should include both SQL data and uploaded files.", "Production deployment should replace default XAMPP credentials with secure database users."])
    return save(doc, "10_APPENDICES.docx")


def index_doc(paths, schema):
    doc = new_doc("PRIMEHR System Documentation Package Index", "Generated professional documentation package")
    add_toc_list(doc, ["Package Contents", "Analysis Summary"])
    doc.add_heading("1. Package Contents", 1)
    add_table(doc, ["No.", "Document", "Filename"], [[str(i), p.stem.replace("_", " ").title(), p.name] for i, p in enumerate(paths, 1)], [.5, 3.0, 3.0])
    doc.add_heading("2. Analysis Summary", 1)
    add_table(doc, ["Item", "Detected Value"], [["PHP files", str(len(list((ROOT/'accounts').rglob('*.php'))))], ["SQL files", str(len(list((ROOT/'database').glob('*.sql'))))], ["Database tables", str(len(schema))], ["Main modules", str(len(MODULES))]], [2.0, 4.5])
    return save(doc, "00_PACKAGE_INDEX.docx")


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    schema = extract_schema()
    (OUT_DIR / "database_schema_summary.json").write_text(json.dumps(schema, indent=2), encoding="utf-8")
    builders = [system_overview, user_manual, admin_guide, technical_doc, database_doc, installation_doc, troubleshooting_doc, security_doc, glossary_doc, appendices_doc]
    paths = [b(schema) for b in builders]
    paths.insert(0, index_doc(paths, schema))
    zip_path = OUT_DIR.parent / "PRIMEHR_System_Documentation_Package.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
        for p in paths + [OUT_DIR / "database_schema_summary.json"]:
            z.write(p, p.name)
    print(zip_path)


if __name__ == "__main__":
    main()
